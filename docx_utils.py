import docx
import logging
from io import BytesIO
from python_redlines.engines import XmlPowerToolsEngine

logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path):
    """
    Extract text content from a DOCX file.
    Returns the full text as a single string.
    """
    try:
        doc = docx.Document(file_path)
        paragraphs_text = [para.text for para in doc.paragraphs if para.text.strip()]
        return '\n'.join(paragraphs_text)
    except Exception as e:
        print(f"Error reading DOCX file {file_path}: {e}")
        raise # Re-raise the exception to be handled upstream


def text_to_docx_bytes(text):
    new_document = docx.Document()
    new_document.add_paragraph(text)
    docx_bytes_io = BytesIO()
    new_document.save(docx_bytes_io)
    return docx_bytes_io.getvalue()


def create_diff_docx(original_path, modified_bytes, output_path):
    docx_comparer = XmlPowerToolsEngine()
    diff_docx_obj = docx_comparer.run_redline(
        author_tag='FeedbackLM', 
        original=original_path, # path or bytes
        modified=modified_bytes) # path or bytes
    diff_docx_bytes = diff_docx_obj[0]
    revision_count = diff_docx_obj[1].strip().split(" ")[-1]
    with open(output_path, 'wb') as f:
        f.write(diff_docx_bytes)
    return revision_count


if __name__ == "__main__":
    print("Running docx_utils directly for testing...")
    original_docx_path = 'examples/original_english.docx'
    improved_docx_path = 'examples/improved_text.docx' # create it beforehand with your LLM of choice
    print("Extracting text from docx")
    original_text = extract_text_from_docx(original_docx_path)
    improved_text = extract_text_from_docx(improved_docx_path)
    print("Converting text to bytes")
    original_docx_bytes = text_to_docx_bytes(original_text)
    improved_docx_bytes = text_to_docx_bytes(improved_text)
    output_path = "examples/test_changes.docx"
    print("Creating document with track changes on")
    revision_count = create_diff_docx(original_docx_bytes, improved_docx_bytes, output_path)

    print("Direct test finished.")